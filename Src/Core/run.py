# coding=utf-8
import sys
import time
import os
import re

import docx
import xlwings
import logging
from Src.Comm.comm import get_cf_value
from docx import Document

logger = logging.getLogger()

scores_dict = {}


def run():
    try:
        scores_calc = get_cf_value("Output", "scores")

        app = xlwings.App(visible=False, add_book=False)
        app.display_alerts = False
        app.screen_updating = False

        wb_scores = app.books.open(scores_calc)
        sht_score = wb_scores.sheets["Score"]
        sht_ans_detail = wb_scores.sheets["Ans Detail"]
        if sht_score.used_range.last_cell.row > 1:
            sht_score['A2:E' + str(sht_score.used_range.last_cell.row)].clear()
        if sht_ans_detail.used_range.last_cell.row > 1:
            sht_ans_detail['A2:G' + str(sht_ans_detail.used_range.last_cell.row)].clear()

        quiz_files_dir = get_cf_value("Input", "quizFilesDir")

        files = os.listdir(quiz_files_dir)

        for f in files:
            file_type = f.split(".")[-1].lower()
            if file_type == "docx":
                get_ans_detail(os.path.join(quiz_files_dir, f), wb_scores)

        wb_scores.save()
        wb_scores.close()
        app.quit()
    except Exception as e:
        logger.info("run failed:" + str(e))
    else:
        logger.info("run end")


def get_ans_detail(f_docx, wb_scores):
    try:
        print("Reading:" + f_docx)
        sht_score = wb_scores.sheets["Score"]
        sht_ans_detail = wb_scores.sheets["Ans Detail"]
        sht_standard_ans = wb_scores.sheets["Standard Ans"]

        doc = docx.Document(f_docx)

        na = search_doc_value(doc, r"姓名[:\s_]*(.*[^_\s])", 0)
        dept = search_doc_value(doc, r"部門[:\s_]*(.*[^ _\s])", 0)
        pos = search_doc_value(doc, r"職位[:\s_]*(.*[^\s_])", 0)
        print(na, dept, pos)

        for i in range(2, sht_standard_ans.used_range.last_cell.row + 1):
            ques_id = sht_standard_ans.range("A" + str(i)).value
            standard_ans = sht_standard_ans.range("B" + str(i)).value
            math_model = sht_standard_ans.range("C" + str(i)).value
            ans = search_doc_value(doc, math_model, 0)
            print(ques_id, ans)

            last_row = sht_ans_detail.used_range.last_cell.row
            sht_ans_detail.range("A" + str(1 + last_row)).value = na
            sht_ans_detail.range("B" + str(1 + last_row)).value = dept
            sht_ans_detail.range("C" + str(1 + last_row)).value = pos
            sht_ans_detail.range("D" + str(1 + last_row)).value = ques_id
            sht_ans_detail.range("E" + str(1 + last_row)).value = ans
            sht_ans_detail.range("F" + str(1 + last_row)).value = standard_ans
            sht_ans_detail.range("G" + str(1 + last_row)).formula = \
                "=IF(E" + str(1 + last_row) + "=F" + str(1 + last_row) + ",1,0)"

        last_row = sht_score.used_range.last_cell.row
        sht_score.range("A" + str(1 + last_row)).value = na
        sht_score.range("B" + str(1 + last_row)).value = dept
        sht_score.range("C" + str(1 + last_row)).value = pos
        sht_score.range("D" + str(1 + last_row)).formula = \
            "=SUMIFS('Ans Detail'!G:G,'Ans Detail'!A:A,Score!A" + str(1 + last_row) \
            + ",'Ans Detail'!B:B,Score!B" + str(1 + last_row) + ")"

        if f_docx[0:1] == ".":
            sht_score.range("E" + str(1 + last_row)).value = os.path.abspath('.') + f_docx[1:]
        else:
            sht_score.range("E" + str(1 + last_row)).value = f_docx

    except Exception as e:
        logger.info("get_ans_detail failed:" + f_docx + " " + str(e))
    else:
        logger.info("get_ans_detail end:" + f_docx)


def search_doc_value(doc, reg_exp, i_start_paragraph):

    for i in range(len(doc.paragraphs)):

        if i < i_start_paragraph:
            continue
        ret = re.search(reg_exp, doc.paragraphs[i].text)
        if ret is not None:
            return ret.groups()[0]

    return ""



